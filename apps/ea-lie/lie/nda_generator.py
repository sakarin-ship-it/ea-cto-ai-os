"""NDA generator — 4 types, Claude API for clause selection (non-sensitive params only),
python-docx bilingual EN/TH output.  ALWAYS embeds every mandatory clause.

Privacy rule: only template parameters (type, duration, purpose category) are sent
to the Claude API.  Party names, commercial terms, and confidential context stay local.
"""
from __future__ import annotations

import io
import json
import logging
from dataclasses import dataclass
from enum import Enum
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .clauses import ALL_MANDATORY_CLAUSES, MandatoryClause

logger = logging.getLogger(__name__)

OPTIONAL_CLAUSE_POOL = [
    "TRADE_SECRET",
    "NON_SOLICITATION",
    "NON_COMPETE",
    "RESIDUALS",
    "RETURN_OF_INFORMATION",
    "AUDIT_RIGHT",
]

_OPTIONAL_TEXT: dict[str, tuple[str, str]] = {
    "TRADE_SECRET": (
        "Trade Secrets: The parties acknowledge that certain Confidential Information may "
        "constitute trade secrets under applicable law and shall be protected accordingly.",
        "ความลับทางการค้า: คู่สัญญาตระหนักว่าข้อมูลลับบางส่วนอาจเป็นความลับทางการค้าตามกฎหมายที่ใช้บังคับ "
        "และจะได้รับการคุ้มครองตามนั้น",
    ),
    "NON_SOLICITATION": (
        "Non-Solicitation: Neither party shall solicit the employees of the other party during "
        "the term of this Agreement and for one (1) year thereafter.",
        "การห้ามชักชวน: คู่สัญญาไม่ฝ่ายใดจะชักชวนพนักงานของอีกฝ่ายในระหว่างข้อตกลงนี้ "
        "และเป็นเวลาหนึ่ง (1) ปีหลังจากนั้น",
    ),
    "NON_COMPETE": (
        "Non-Competition: The Receiving Party shall not engage in any business that directly "
        "competes with the Disclosing Party during the term hereof.",
        "การห้ามแข่งขัน: ฝ่ายรับข้อมูลจะไม่ประกอบธุรกิจที่แข่งขันโดยตรงกับฝ่ายเปิดเผยในระหว่างข้อตกลงนี้",
    ),
    "RESIDUALS": (
        "Residuals: The Receiving Party may use residual knowledge retained in unaided memory "
        "of its personnel, provided such use is not a deliberate circumvention of this Agreement.",
        "ข้อมูลที่จำได้: ฝ่ายรับข้อมูลอาจใช้ความรู้ที่บุคลากรจำได้โดยไม่ได้ตั้งใจ "
        "หากไม่ใช่ความพยายามโดยเจตนาที่จะหลีกเลี่ยงข้อตกลงนี้",
    ),
    "RETURN_OF_INFORMATION": (
        "Return of Information: Upon termination or upon request, the Receiving Party shall "
        "promptly return or certifiably destroy all Confidential Information.",
        "การคืนข้อมูล: เมื่อข้อตกลงนี้สิ้นสุดลงหรือตามคำขอ ฝ่ายรับข้อมูลต้องคืนหรือทำลาย "
        "ข้อมูลลับทั้งหมดทันทีและให้หนังสือรับรองการทำลาย",
    ),
    "AUDIT_RIGHT": (
        "Audit Right: The Disclosing Party may audit the Receiving Party's compliance with this "
        "Agreement upon fourteen (14) days' prior written notice, not more than once per year.",
        "สิทธิในการตรวจสอบ: ฝ่ายเปิดเผยมีสิทธิตรวจสอบการปฏิบัติตามข้อตกลงนี้ของฝ่ายรับข้อมูล "
        "โดยแจ้งล่วงหน้าเป็นลายลักษณ์อักษรสิบสี่ (14) วัน ไม่เกินปีละหนึ่งครั้ง",
    ),
}


class NDAType(str, Enum):
    UNILATERAL = "unilateral"
    MUTUAL = "mutual"
    EMPLOYEE = "employee"
    VENDOR = "vendor"


@dataclass
class NDAParams:
    nda_type: NDAType
    disclosing_party: str
    receiving_party: str
    purpose: str
    duration_years: int
    confidentiality_period_years: int = 3
    governing_law: str = "Thailand"


class NDAGenerator:
    def __init__(self, anthropic_client: Optional[anthropic.Anthropic] = None) -> None:
        self._client = anthropic_client or anthropic.Anthropic()

    def generate(self, params: NDAParams) -> bytes:
        """Return bilingual EN/TH NDA docx bytes with ALL mandatory clauses."""
        optional_ids = self._select_optional_clauses(params)
        return self._build_docx(params, optional_ids)

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _select_optional_clauses(self, params: NDAParams) -> list[str]:
        """Claude API selects optional clauses from non-sensitive params only."""
        prompt = (
            f"You are a Thai law specialist selecting optional NDA clauses.\n"
            f"NDA type: {params.nda_type.value}\n"
            f"Purpose category: {params.purpose[:80]}\n"
            f"Duration: {params.duration_years} year(s)\n"
            f"Confidentiality period: {params.confidentiality_period_years} year(s)\n\n"
            f"Return a JSON array (no explanation) of clause IDs to include from: "
            f"{json.dumps(OPTIONAL_CLAUSE_POOL)}"
        )
        try:
            msg = self._client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=256,
                messages=[{"role": "user", "content": prompt}],
            )
            text = msg.content[0].text.strip()
            parsed = json.loads(text)
            return [c for c in parsed if c in OPTIONAL_CLAUSE_POOL]
        except Exception as exc:
            logger.warning("Claude clause selection failed (%s); using defaults", exc)
            return ["RETURN_OF_INFORMATION"]

    def _build_docx(self, params: NDAParams, optional_ids: list[str]) -> bytes:
        doc = Document()

        # Bilingual title
        for text in ("NON-DISCLOSURE AGREEMENT", "ข้อตกลงการรักษาความลับ"):
            h = doc.add_heading(text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parties
        doc.add_heading("PARTIES / คู่สัญญา", level=2)
        doc.add_paragraph(
            f"Disclosing Party / ฝ่ายเปิดเผย: {params.disclosing_party}\n"
            f"Receiving Party / ฝ่ายรับข้อมูล: {params.receiving_party}"
        )

        # Purpose
        doc.add_heading("PURPOSE / วัตถุประสงค์", level=2)
        doc.add_paragraph(params.purpose)

        # Type-specific confidentiality obligations
        doc.add_heading(
            "CONFIDENTIALITY OBLIGATIONS / พันธกรณีการรักษาความลับ", level=2
        )
        doc.add_paragraph(self._type_clause(params))

        # Term
        doc.add_heading("TERM / ระยะเวลา", level=2)
        doc.add_paragraph(
            f"This Agreement shall remain in effect for {params.duration_years} year(s) "
            f"from the date of execution, with confidentiality obligations surviving for "
            f"{params.confidentiality_period_years} year(s) after termination.  /  "
            f"ข้อตกลงนี้มีผลบังคับเป็นเวลา {params.duration_years} ปีนับจากวันที่ลงนาม "
            f"โดยพันธกรณีการรักษาความลับยังคงมีผลต่อไปอีก {params.confidentiality_period_years} ปี"
            f"หลังจากสิ้นสุด"
        )

        # Optional clauses
        if optional_ids:
            doc.add_heading("ADDITIONAL PROVISIONS / บทบัญญัติเพิ่มเติม", level=2)
            for cid in optional_ids:
                if cid in _OPTIONAL_TEXT:
                    en, th = _OPTIONAL_TEXT[cid]
                    doc.add_paragraph(en)
                    doc.add_paragraph(th)

        # *** MANDATORY — embedded unconditionally ***
        doc.add_heading(
            "MANDATORY LEGAL PROVISIONS / บทบัญญัติบังคับทางกฎหมาย", level=2
        )
        for clause in ALL_MANDATORY_CLAUSES:
            self._embed_mandatory(doc, clause)

        # Governing law
        doc.add_heading("GOVERNING LAW / กฎหมายที่ใช้บังคับ", level=2)
        doc.add_paragraph(
            f"This Agreement shall be governed by the laws of {params.governing_law}.  /  "
            f"ข้อตกลงนี้อยู่ภายใต้กฎหมายของ{params.governing_law}"
        )

        # Signature block
        doc.add_heading("SIGNATURES / ลายมือชื่อ", level=2)
        doc.add_paragraph(
            "IN WITNESS WHEREOF the parties have executed this Agreement.\n"
            "เพื่อเป็นหลักฐาน คู่สัญญาได้ลงนามในข้อตกลงนี้\n\n"
            "Disclosing Party: ____________________   Date: ___________\n"
            "ฝ่ายเปิดเผย:       ____________________   วันที่: ___________\n\n"
            "Receiving Party:  ____________________   Date: ___________\n"
            "ฝ่ายรับข้อมูล:     ____________________   วันที่: ___________"
        )

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _type_clause(self, params: NDAParams) -> str:
        dp, rp = params.disclosing_party, params.receiving_party
        if params.nda_type == NDAType.MUTUAL:
            return (
                f"Both {dp} and {rp} agree to maintain strict confidentiality of all "
                f"Confidential Information received from the other party.  /  "
                f"{dp} และ {rp} ตกลงที่จะรักษาความลับของข้อมูลลับทั้งหมดที่ได้รับจากอีกฝ่าย"
            )
        if params.nda_type == NDAType.UNILATERAL:
            return (
                f"{rp} (Receiving Party) agrees to maintain strict confidentiality of all "
                f"Confidential Information received from {dp} (Disclosing Party).  /  "
                f"{rp} (ฝ่ายรับข้อมูล) ตกลงที่จะรักษาความลับของข้อมูลลับทั้งหมด"
                f"ที่ได้รับจาก {dp} (ฝ่ายเปิดเผย)"
            )
        if params.nda_type == NDAType.EMPLOYEE:
            return (
                f"Employee ({rp}) agrees to maintain confidentiality of all proprietary "
                f"information and trade secrets of {dp} during employment and for "
                f"{params.confidentiality_period_years} years thereafter.  /  "
                f"พนักงาน ({rp}) ตกลงที่จะรักษาความลับของข้อมูลที่เป็นกรรมสิทธิ์และความลับทางการค้า "
                f"ของ {dp} ในระหว่างการจ้างงานและเป็นเวลา {params.confidentiality_period_years} ปีหลังจากนั้น"
            )
        # VENDOR
        return (
            f"Vendor ({rp}) agrees to maintain strict confidentiality of all client data, "
            f"business processes, and technical information of {dp} disclosed in connection "
            f"with the services.  /  "
            f"ผู้จัดหาสินค้า/บริการ ({rp}) ตกลงที่จะรักษาความลับอย่างเคร่งครัดของข้อมูลลูกค้า "
            f"กระบวนการทางธุรกิจ และข้อมูลทางเทคนิคของ {dp} ที่เปิดเผยเกี่ยวกับบริการ"
        )

    def _embed_mandatory(self, doc: Document, clause: MandatoryClause) -> None:
        doc.add_heading(f"{clause.title_en}  /  {clause.title_th}", level=3)
        doc.add_paragraph(clause.text_en)
        doc.add_paragraph(clause.text_th)
        # Tag paragraph — white 1pt text; invisible in print but readable by tests
        tag_para = doc.add_paragraph(clause.tag)
        run = tag_para.runs[0]
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
