"""Redline generator — python-docx tracked-style changes + mandatory watermark.

Watermark "AI DRAFT — LEGAL COUNSEL REVIEW REQUIRED" is added to the header of
every section.  Tests assert its presence unconditionally.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

WATERMARK_TEXT = "AI DRAFT — LEGAL COUNSEL REVIEW REQUIRED"


@dataclass
class RedlineChange:
    original_text: str
    revised_text: str
    reason: str = ""
    clause_ref: str = ""


class RedlineGenerator:
    """Generate redlined docx with watermark from an original document."""

    def __init__(self, watermark_text: str = WATERMARK_TEXT) -> None:
        self._watermark = watermark_text

    @property
    def watermark_text(self) -> str:
        return self._watermark

    def generate(self, original_docx: bytes, changes: list[RedlineChange]) -> bytes:
        """Return a new docx bytes with watermark + change markings applied."""
        doc = Document(io.BytesIO(original_docx))
        self._apply_watermark(doc)
        self._apply_changes(doc, changes)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------

    def _apply_watermark(self, doc: Document) -> None:
        """Add watermark to every section header — ALWAYS."""
        for section in doc.sections:
            header = section.header
            if not header.paragraphs:
                header.add_paragraph()
            para = header.paragraphs[0]
            para.clear()
            run = para.add_run(self._watermark)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_changes(self, doc: Document, changes: list[RedlineChange]) -> None:
        for para in doc.paragraphs:
            for change in changes:
                if change.original_text and change.original_text in para.text:
                    self._mark_deletion(para, change.original_text)
                    self._mark_insertion(para, change.revised_text)

    @staticmethod
    def _mark_deletion(para, text: str) -> None:
        for run in para.runs:
            if text in run.text:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.font.strike = True
                break

    @staticmethod
    def _mark_insertion(para, text: str) -> None:
        run = para.add_run(f" {text}")
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        run.font.underline = True
