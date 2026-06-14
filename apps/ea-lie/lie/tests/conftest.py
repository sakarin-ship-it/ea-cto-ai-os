"""Shared pytest fixtures for EA-LIE tests."""
from __future__ import annotations

import io

import pytest
from docx import Document


@pytest.fixture
def blank_docx_bytes() -> bytes:
    """Return bytes of an empty docx document."""
    doc = Document()
    doc.add_paragraph("Original clause text for testing.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extract_all_text(docx_bytes: bytes) -> str:
    """Extract all paragraph text from a docx, including invisible runs."""
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_header_text(docx_bytes: bytes) -> str:
    """Extract all header paragraph text from a docx."""
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            lines.append(p.text)
    return "\n".join(lines)
