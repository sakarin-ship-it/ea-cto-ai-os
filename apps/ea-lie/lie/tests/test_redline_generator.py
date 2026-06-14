"""Tests for redline generator — watermark ALWAYS present."""
from __future__ import annotations

import io

import pytest
from docx import Document

from lie.redline_generator import WATERMARK_TEXT, RedlineChange, RedlineGenerator
from lie.tests.conftest import extract_header_text


@pytest.fixture
def generator():
    return RedlineGenerator()


@pytest.fixture
def simple_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("This agreement shall be for a period of one (1) year.")
    doc.add_paragraph("The fee shall be THB 500,000 per month.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# CRITICAL: watermark always present
# ---------------------------------------------------------------------------

def test_watermark_present_no_changes(generator, simple_docx):
    """Watermark must be in header even with zero changes."""
    result = generator.generate(simple_docx, [])
    header = extract_header_text(result)
    assert WATERMARK_TEXT in header, (
        f"Watermark '{WATERMARK_TEXT}' not found in document header"
    )


def test_watermark_present_with_changes(generator, simple_docx):
    """Watermark survives redline application."""
    changes = [
        RedlineChange(
            original_text="one (1) year",
            revised_text="two (2) years",
            reason="Term extension per Amendment 1",
        )
    ]
    result = generator.generate(simple_docx, changes)
    header = extract_header_text(result)
    assert WATERMARK_TEXT in header


def test_watermark_custom_text():
    """Custom watermark text is honoured."""
    custom = "DRAFT — DO NOT DISTRIBUTE"
    gen = RedlineGenerator(watermark_text=custom)
    doc = Document()
    doc.add_paragraph("Test")
    buf = io.BytesIO()
    doc.save(buf)
    result = gen.generate(buf.getvalue(), [])
    header = extract_header_text(result)
    assert custom in header


def test_default_watermark_constant():
    assert WATERMARK_TEXT == "AI DRAFT — LEGAL COUNSEL REVIEW REQUIRED"


# ---------------------------------------------------------------------------
# Change tracking
# ---------------------------------------------------------------------------

def test_deletion_marked_strikethrough(generator, simple_docx):
    changes = [
        RedlineChange(
            original_text="one (1) year",
            revised_text="two (2) years",
        )
    ]
    result_bytes = generator.generate(simple_docx, changes)
    doc = Document(io.BytesIO(result_bytes))

    # Find a run with strikethrough
    struck = [
        run
        for para in doc.paragraphs
        for run in para.runs
        if run.font.strike
    ]
    assert struck, "Expected at least one struck-through run for deleted text"


def test_insertion_marked_underline(generator, simple_docx):
    changes = [
        RedlineChange(
            original_text="one (1) year",
            revised_text="two (2) years",
        )
    ]
    result_bytes = generator.generate(simple_docx, changes)
    doc = Document(io.BytesIO(result_bytes))

    underlined = [
        run
        for para in doc.paragraphs
        for run in para.runs
        if run.font.underline
    ]
    assert underlined, "Expected at least one underlined run for inserted text"


def test_output_is_valid_docx(generator, simple_docx):
    result = generator.generate(simple_docx, [])
    assert isinstance(result, bytes)
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) >= 1
