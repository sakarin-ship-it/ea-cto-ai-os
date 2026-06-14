"""RFx generator — Claude API produces bilingual DOCX/PDF from non-sensitive scope."""
from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

import anthropic
from docx import Document
from fpdf import FPDF
from sqlalchemy.orm import Session

from ea_pip.constants import MAX_SCOPE_LENGTH
from ea_pip.models import Package, append_audit


@dataclass
class RFxOutput:
    docx_path: str
    pdf_path: str
    content_en: str
    content_th: str


def _generate_bilingual_content(
    scope_en: str,
    template: str,
    title_en: str,
    title_th: str,
) -> tuple[str, str]:
    """Call Claude API to produce bilingual RFx body. Scope must be non-sensitive."""
    client = anthropic.Anthropic()
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=(
            "You are a bilingual procurement specialist. "
            "Generate formal Request for Proposal (RFP) content in English and Thai. "
            "Return ONLY a JSON object with keys 'english' and 'thai'."
        ),
        messages=[
            {
                "role": "user",
                "content": (
                    f"Package: {title_en} / {title_th}\n\n"
                    f"Template type: {template}\n\n"
                    f"Scope:\n{scope_en}\n\n"
                    "Generate a complete bilingual RFP body. "
                    "Return JSON: {\"english\": \"...\", \"thai\": \"...\"}"
                ),
            }
        ],
    )
    raw = message.content[0].text.strip()
    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
    data = json.loads(raw)
    return data["english"], data["thai"]


def _write_docx(
    title_en: str,
    title_th: str,
    content_en: str,
    content_th: str,
    output_path: str,
) -> None:
    doc = Document()
    doc.add_heading(f"{title_en}", 0)
    doc.add_heading(f"{title_th}", 1)
    doc.add_heading("English", 2)
    for para in content_en.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_page_break()
    doc.add_heading("ภาษาไทย (Thai)", 2)
    for para in content_th.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save(output_path)


def _write_pdf(
    title_en: str,
    content_en: str,
    content_th: str,
    output_path: str,
) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title_en[:120])
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "English", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 6, content_en[:5000])
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Thai (see DOCX for full Thai text)", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)
    # fpdf2 core fonts lack Thai glyphs; write ASCII transliteration stub
    pdf.multi_cell(0, 6, "[Thai content — see accompanying .docx for full bilingual version]")
    pdf.output(output_path)


def generate_rfx(
    package_id: int,
    template: str,
    output_dir: str,
    actor: str,
    session: Session,
) -> RFxOutput:
    """Generate bilingual RFx DOCX + PDF for a package. Scope must be non-sensitive."""
    package = session.get(Package, package_id)
    if package is None:
        raise ValueError(f"Package {package_id} not found")

    if len(package.scope_en) > MAX_SCOPE_LENGTH:
        raise ValueError(
            f"scope_en length {len(package.scope_en)} exceeds MAX_SCOPE_LENGTH={MAX_SCOPE_LENGTH}. "
            "Only non-sensitive scope briefs may be sent to the Claude API (CLAUDE.md rule 1)."
        )

    base_dir = Path(output_dir).resolve()
    base_dir.mkdir(parents=True, exist_ok=True)
    base = f"rfx_{package.package_no}"
    docx_path = str(base_dir / f"{base}.docx")
    pdf_path = str(base_dir / f"{base}.pdf")

    # Defense-in-depth: package_no validated by API at creation, but verify no traversal
    for _p in (Path(docx_path), Path(pdf_path)):
        if not _p.resolve().is_relative_to(base_dir):
            raise ValueError(f"Resolved output path escapes output_dir: {_p}")

    content_en, content_th = _generate_bilingual_content(
        scope_en=package.scope_en,
        template=template,
        title_en=package.title_en,
        title_th=package.title_th,
    )

    _write_docx(package.title_en, package.title_th, content_en, content_th, docx_path)
    _write_pdf(package.title_en, content_en, content_th, pdf_path)

    package.rfx_docx_path = docx_path
    package.rfx_pdf_path = pdf_path

    append_audit(
        session,
        entity_type="package",
        entity_id=package_id,
        action="rfx_generated",
        actor=actor,
        payload={"docx_path": docx_path, "pdf_path": pdf_path, "template": template},
    )

    return RFxOutput(
        docx_path=docx_path,
        pdf_path=pdf_path,
        content_en=content_en,
        content_th=content_th,
    )
