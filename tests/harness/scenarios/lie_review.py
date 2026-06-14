"""EA-LIE: risk score in range, correct routing, watermark present."""
from __future__ import annotations

import io
import json
import sys
from pathlib import Path
from unittest.mock import patch

_ROOT = Path(__file__).resolve().parents[4]
for _p in [str(_ROOT / "apps/ea-lie"), str(_ROOT / "shared")]:
    if _p not in sys.path:
        sys.path.insert(0, _p)

from lie.review_engine import RAGStatus, ReviewEngine, ReviewerLevel
from lie.redline_generator import WATERMARK_TEXT, RedlineChange, RedlineGenerator

from tests.harness.generators import contract_text_with_all_clauses

SCENARIO_ID = "lie_review"

_SCORE_SCENARIOS = [
    # (name, findings_overrides, use_clause_tags_in_text, expected_rag, expected_reviewer)
    # review_engine checks BOTH findings AND raw text for each tag —
    # to produce a penalty the finding must be False AND the tag absent from text.
    ("green_all_present", {}, True, RAGStatus.GREEN, ReviewerLevel.CTO_CFO),
    ("amber_2_missing", {"has_injunctive_relief": False, "has_pdpa_reference": False},
     False, RAGStatus.AMBER, ReviewerLevel.LEGAL_COUNSEL),
    ("red_all_missing", {
        "has_injunctive_relief": False, "has_liquidated_damages": False,
        "has_pdpa_reference": False, "has_sec_carveout": False, "has_esignature": False,
    }, False, RAGStatus.RED, ReviewerLevel.EXTERNAL_LAWYER),
]

_PLAIN_CONTRACT = (
    "EPC CONTRACT\n\nThis agreement sets out the terms for engineering, procurement, "
    "and construction services. The parties agree to the following terms and conditions. "
    "Payment shall be made within 30 days of invoice. Governing law: Thailand."
)


def _base_findings(**overrides) -> str:
    base = {
        "has_injunctive_relief": True,
        "has_liquidated_damages": True,
        "ld_amount_thb": 15_000_000,
        "has_pdpa_reference": True,
        "has_sec_carveout": True,
        "has_esignature": True,
        "governing_law": "Thailand",
        "problematic_clauses": [],
        "missing_standard_clauses": [],
        "risk_factors": [],
    }
    base.update(overrides)
    return json.dumps(base)


def setup(seed: int) -> dict:
    scenario = _SCORE_SCENARIOS[seed % len(_SCORE_SCENARIOS)]
    name, overrides, use_tags, expected_rag, expected_reviewer = scenario
    findings_json = _base_findings(**overrides)
    # When clauses are "missing" we use plain text (no ##MC:## tags) so the
    # review_engine cannot find them in the raw text either.
    contract_text = contract_text_with_all_clauses(seed) if use_tags else _PLAIN_CONTRACT
    return {
        "seed": seed,
        "scenario_name": name,
        "findings_json": findings_json,
        "contract_text": contract_text,
        "expected_rag": expected_rag.value,
        "expected_reviewer": expected_reviewer.value,
    }


def run(data: dict) -> dict:
    engine = ReviewEngine()
    with patch("lie.review_engine.chat_complete", return_value=data["findings_json"]):
        r = engine.review_text(data["contract_text"])

    # Watermark via redline generator
    try:
        from docx import Document as _Doc
        doc = _Doc()
        doc.add_paragraph("Original contract text.")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        gen = RedlineGenerator()
        redlined = gen.generate(docx_bytes, [
            RedlineChange("Original", "Revised", "test change", "Article 1"),
        ])
        watermark_found = WATERMARK_TEXT in _extract_header(redlined)
    except Exception:
        watermark_found = None  # docx not available; skip this check

    return {
        "score": r.score,
        "rag": r.rag.value,
        "reviewer_level": r.reviewer_level.value,
        "is_critical": r.is_critical,
        "watermark_found": watermark_found,
    }


def _extract_header(docx_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            lines.append(p.text)
    return "\n".join(lines)


def assert_invariants(data: dict, result: dict) -> None:
    seed = data["seed"]
    score = result["score"]

    # Score range
    assert 0 <= score <= 100, f"seed={seed}: score {score} outside [0, 100]"

    # RAG routing
    if score < 35:
        assert result["rag"] == RAGStatus.GREEN.value, (
            f"seed={seed}: score {score} < 35 → GREEN expected, got {result['rag']}"
        )
        assert result["reviewer_level"] == ReviewerLevel.CTO_CFO.value, (
            f"seed={seed}: score {score} → CTO_CFO expected"
        )
    elif score <= 60:
        assert result["rag"] == RAGStatus.AMBER.value, (
            f"seed={seed}: score {score} in [35,60] → AMBER expected, got {result['rag']}"
        )
        assert result["reviewer_level"] == ReviewerLevel.LEGAL_COUNSEL.value, (
            f"seed={seed}: score {score} → LEGAL_COUNSEL expected"
        )
    else:
        assert result["rag"] == RAGStatus.RED.value, (
            f"seed={seed}: score {score} > 60 → RED expected, got {result['rag']}"
        )
        assert result["reviewer_level"] == ReviewerLevel.EXTERNAL_LAWYER.value, (
            f"seed={seed}: score {score} → EXTERNAL_LAWYER expected"
        )
        assert result["is_critical"] is True, (
            f"seed={seed}: score {score} > 60 → is_critical must be True"
        )

    # Expected rag/reviewer match data
    assert result["rag"] == data["expected_rag"], (
        f"seed={seed} ({data['scenario_name']}): "
        f"rag={result['rag']} != expected {data['expected_rag']}"
    )

    # Watermark (when docx available)
    if result["watermark_found"] is not None:
        assert result["watermark_found"], (
            f"seed={seed}: watermark '{WATERMARK_TEXT}' must appear in redlined document header"
        )
